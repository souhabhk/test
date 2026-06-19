from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Any

from docx import Document
from docx.text.paragraph import Paragraph

@dataclass
class SectionValidationResult:
    section: str
    is_present: bool
    is_mandatory: bool
    min_words: int
    actual_words: int
    passed: bool
    errors: List[str] = field(default_factory=list)

@dataclass
class ValidationReport:
    is_valid: bool = True
    section_results: List[SectionValidationResult] = field(default_factory=list)
    total_words: int = 0
    global_errors: List[str] = field(default_factory=list)
    # SELF-HEAL: Added missing fields used by subclass linters to prevent AttributeError
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)

class BaseDocumentLinter:
    """Classe de base pour la validation déterministe de documents .docx."""
    
    def __init__(self, template_path: Optional[str] = None):
        # SELF-HEAL: Added template_path parameter to match subclass initialization and plan intent
        self.template_rules: Dict[str, Dict[str, Any]] = {}
        self._total_rule: Optional[Dict[str, Any]] = None
        if template_path:
            self.load_template(template_path)

    @staticmethod
    def _validate_path(path: str) -> Path:
        """Résout et valide le chemin pour prévenir les attaques par traversal."""
        resolved = Path(path).resolve()
        if not resolved.exists():
            raise FileNotFoundError(f"Fichier introuvable : {path}")
        if resolved.suffix.lower() != '.docx':
            raise ValueError("Seuls les fichiers .docx sont autorisés.")
        return resolved

    def load_template(self, template_path: str) -> Dict[str, Dict[str, Any]]:
        """Charge et parse le tableau de règles depuis un template .docx."""
        resolved_path = self._validate_path(template_path)
        doc = Document(str(resolved_path))
        rules: Dict[str, Dict[str, Any]] = {}
        
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) < 3:
                    continue
                    
                section_name = cells[0]
                mandatory_str = cells[1].lower()
                try:
                    min_words = int(cells[2])
                except ValueError:
                    min_words = 0
                    
                is_mandatory = mandatory_str in ("oui", "yes", "true", "1")
                
                if section_name.lower() == "total document":
                    self._total_rule = {"mandatory": True, "min_words": min_words}
                else:
                    rules[section_name] = {"mandatory": is_mandatory, "min_words": min_words}
                    
        self.template_rules = rules
        return rules

    def extract_sections(self, doc_path: str) -> Dict[str, List[Paragraph]]:
        """Extrait les paragraphes groupés par titres de niveau 1."""
        resolved_path = self._validate_path(doc_path)
        doc = Document(str(resolved_path))
        sections: Dict[str, List[Paragraph]] = {}
        current_section: Optional[str] = None
        
        for para in doc.paragraphs:
            if para.style.name == 'Heading 1':
                current_section = para.text.strip()
                # SELF-HEAL: Prevent empty heading text from creating a dummy section key that swallows subsequent paragraphs
                if current_section and current_section not in sections:
                    sections[current_section] = []
            elif current_section:
                sections[current_section].append(para)
                
        return sections

    def count_words(self, paragraphs: List[Paragraph]) -> int:
        """Compte les mots dans une liste de paragraphes, ignorant les espaces vides."""
        total = 0
        for para in paragraphs:
            text = para.text.strip()
            if text:
                total += len(text.split())
        return total

    def validate_structure(self, target_doc_path: str) -> ValidationReport:
        """Valide la structure et les seuils de mots selon les règles du template."""
        report = ValidationReport()
        sections = self.extract_sections(target_doc_path)
        
        all_words = 0
        
        for section_name, rule in self.template_rules.items():
            res = SectionValidationResult(
                section=section_name,
                is_mandatory=rule["mandatory"],
                min_words=rule["min_words"],
                is_present=False,
                actual_words=0,
                passed=True
            )
            
            if section_name in sections:
                res.is_present = True
                res.actual_words = self.count_words(sections[section_name])
                all_words += res.actual_words
                
                if res.actual_words < res.min_words:
                    res.passed = False
                    res.errors.append(f"Moins de mots requis ({res.actual_words}/{res.min_words})")
            else:
                if res.is_mandatory:
                    res.passed = False
                    res.errors.append("Section obligatoire manquante")
                    
            report.section_results.append(res)
            if not res.passed:
                report.is_valid = False
                
        if self._total_rule:
            total_min = self._total_rule["min_words"]
            if all_words < total_min:
                report.is_valid = False
                report.global_errors.append(f"Total mots insuffisant ({all_words}/{total_min})")
                
        report.total_words = all_words
        return report