import os
from docx import Document


def _create_template(
    output_path: str,
    headings: list[str],
    rules: list[tuple[str, str, int]]
) -> None:
    """Generates a .docx reference template with specified headings and a validation rules table."""
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = Document()
    
    doc.add_heading("Template de Validation", level=0)
    doc.add_paragraph(
        "Ce document définit la structure attendue et les règles de conformité "
        "pour la validation automatique des documents générés."
    )
    
    for heading in headings:
        doc.add_heading(heading, level=1)
        
    doc.add_heading("Règles de conformité", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Section"
    hdr_cells[1].text = "Obligatoire"
    hdr_cells[2].text = "Mots minimum"
    
    for section_name, is_mandatory, min_words in rules:
        row_cells = table.add_row().cells
        row_cells[0].text = section_name
        row_cells[1].text = is_mandatory
        row_cells[2].text = str(min_words)
        
    doc.save(output_path)
    print(f"[OK] Template généré : {output_path}")


def create_spec_template() -> None:
    """Generates the SPEC reference template."""
    headings = [
        "Introduction",
        "Contexte",
        "Acteurs",
        "Exigences fonctionnelles",
        "Contraintes",
        "Glossaire"
    ]
    rules = [
        ("Introduction", "Oui", 50),
        ("Contexte", "Oui", 100),
        ("Acteurs", "Oui", 30),
        ("Exigences fonctionnelles", "Oui", 200),
        ("Contraintes", "Oui", 50),
        ("Glossaire", "Non", 20),
        ("Total document", "Oui", 500)
    ]
    _create_template("templates/spec_template.docx", headings, rules)


def create_archi_template() -> None:
    """Generates the Architecture reference template."""
    headings = [
        "Vue d'ensemble",
        "Composants",
        "Flux de données",
        "Déploiement",
        "Sécurité",
        "Décisions techniques (ADR)"
    ]
    rules = [
        ("Vue d'ensemble", "Oui", 80),
        ("Composants", "Oui", 150),
        ("Flux de données", "Oui", 100),
        ("Déploiement", "Oui", 50),
        ("Sécurité", "Oui", 40),
        ("Décisions techniques", "Non", 50),
        ("Total document", "Oui", 600)
    ]
    _create_template("templates/archi_template.docx", headings, rules)


if __name__ == "__main__":
    print("Génération des templates de référence...")
    create_spec_template()
    create_archi_template()
    print("Terminé.")