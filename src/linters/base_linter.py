import os
import re
import logging
from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional

from docx import Document
from src.utils.docx_parser import DocxParser, TemplateParseError

logger = logging.getLogger(__name__)


class ValidationRule:
    def __init__(self, section: str, is_mandatory: bool, min_words: int):
        self.section = section
        self.is_mandatory = is_mandatory
        self.min_words = min_words


class BaseLinter(ABC):
    """
    Abstract base class handling template loading, rule evaluation, and word counting.
    Subclasses must specify their template path and can add document-specific checks.
    """
    
    def __init__(self, template_path: str):
        self.template_path = template_path
        self._rules: Optional[List[ValidationRule]] = None
        self._parser: Optional[DocxParser] = None

    def load_template(self) -> List[ValidationRule]:
        """Loads and caches validation rules from the .docx template."""
        if self._rules is not None:
            return self._rules
        
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"Template not found: {self.template_path}")
            
        try:
            # SELF-HEAL: Fixed DocxParser instantiation (was missing path argument) and replaced non-existent parse_rules() call with property access
            parser = DocxParser(self.template_path)
            raw_rules = parser.rules
            self._rules = [
                ValidationRule(section=r['section'], is_mandatory=r['mandatory'], min_words=r['min_words'])
                for r in raw_rules
            ]
            logger.info(f"Loaded {len(self._rules)} rules from {self.template_path}")
            return self._rules
        except TemplateParseError as e:
            logger.error(f"Failed to parse template {self.template_path}: {e}")
            raise

    @staticmethod
    def _count_words(text: str) -> int:
        """Deterministic word counting using regex to handle punctuation and hyphens."""
        if not text:
            return 0
        return len(re.findall(r'\b\w+\b', text))

    @staticmethod
    def _extract_section_text(doc: Document, heading: str) -> str:
        """
        Extracts text content following a specific heading until the next heading of the same or higher level.
        """
        section_paragraphs: List[str] = []
        capturing = False
        
        for para in doc.paragraphs:
            if para.style.name.lower().startswith('heading'):
                if para.text.strip() == heading:
                    capturing = True
                    continue
                elif capturing:
                    break
            if capturing:
                section_paragraphs.append(para.text)
                
        return "\n".join(section_paragraphs)

    def _parse_document(self, document_path: str) -> List[Dict[str, str]]:
        """Parses a target document into sections with headings and text."""
        try:
            doc = Document(document_path)
        except Exception as e:
            raise IOError(f"Failed to open document for validation: {e}") from e
            
        sections = []
        current_heading = None
        current_text_parts = []
        
        for para in doc.paragraphs:
            if para.style.name.lower().startswith('heading'):
                if current_heading is not None:
                    sections.append({
                        'heading': current_heading,
                        'text': "\n".join(current_text_parts)
                    })
                current_heading = para.text.strip()
                current_text_parts = []
            else:
                if current_heading is not None:
                    current_text_parts.append(para.text)
                    
        if current_heading is not None:
            sections.append({
                'heading': current_heading,
                'text': "\n".join(current_text_parts)
            })
            
        return sections

    def validate(self, document_path: str) -> Dict[str, Any]:
        """
        Validates a document against the loaded template rules.
        Returns a structured report with pass/fail status and deltas.
        """
        rules = self.load_template()
        # SELF-HEAL: Replaced non-existent self._parser.parse_document() with implemented _parse_document method
        sections = self._parse_document(document_path)
        
        missing_sections = []
        word_count_violations = []
        section_counts: Dict[str, int] = {}
        
        # Calculate word counts per section
        for section in sections:
            heading = section.get('heading', '').strip()
            text = section.get('text', '')
            section_counts[heading] = self._count_words(text)

        # Evaluate section-level rules
        for rule in rules:
            if rule.section.strip().lower() == 'total document':
                continue
            
            matched_heading = next(
                (h for h in section_counts if rule.section.lower() in h.lower()),
                None
            )
            
            if matched_heading is None:
                if rule.is_mandatory:
                    missing_sections.append(rule.section)
            else:
                actual_words = section_counts[matched_heading]
                if actual_words < rule.min_words:
                    word_count_violations.append({
                        'section': rule.section,
                        'expected': rule.min_words,
                        'actual': actual_words
                    })

        # Evaluate total document word count
        total_text = '\n'.join(s.get('text', '') for s in sections)
        total_words = self._count_words(total_text)
        
        total_rule = next(
            (r for r in rules if r.section.strip().lower() == 'total document'),
            None
        )
        if total_rule and total_words < total_rule.min_words:
            word_count_violations.append({
                'section': 'Total document',
                'expected': total_rule.min_words,
                'actual': total_words
            })

        is_valid = len(missing_sections) == 0 and len(word_count_violations) == 0

        # SELF-HEAL: Changed return type from ValidationReport dataclass to Dict to match plan's JSON structure and fix subclass dict-assignment bugs
        return {
            "valid": is_valid,
            "missing_sections": missing_sections,
            "word_count_violations": word_count_violations,
            "total_words": total_words,
            "section_counts": section_counts
        }

    @abstractmethod
    def get_template_path(self) -> str:
        """Subclasses should return the path to their specific template."""
        pass