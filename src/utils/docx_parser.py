import re
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Tuple, Union

from docx import Document
from docx.opc.exceptions import PackageNotFoundError


class TemplateParseError(Exception):
    """Raised when a .docx template does not conform to the expected structure."""
    pass


@dataclass(frozen=True)
class ValidationRule:
    section: str
    is_mandatory: bool
    min_words: int


class DocxParser:
    """
    Parses Word documents to extract structural elements for validation.
    Designed to work with templates containing a specific rules table.
    """
    
    REQUIRED_HEADERS = {'section', 'obligatoire', 'mots minimum'}
    HEADING_STYLE_PREFIX = 'Heading'
    
    def __init__(self, doc_path: Union[str, Path]):
        self.doc_path = Path(doc_path)
        if not self.doc_path.exists():
            raise FileNotFoundError(f"Document not found: {self.doc_path}")
        
        try:
            self.doc = Document(self.doc_path)
        except PackageNotFoundError as e:
            raise TemplateParseError(f"Invalid .docx file format: {self.doc_path}") from e

    @staticmethod
    def count_words(text: str) -> int:
        """Deterministic word count using regex to handle punctuation and hyphens."""
        if not text:
            return 0
        return len(re.findall(r'\b\w+\b', text))

    def extract_headings(self) -> List[str]:
        """Extracts an ordered list of all heading texts from the document."""
        headings = []
        for para in self.doc.paragraphs:
            if para.style.name.startswith(self.HEADING_STYLE_PREFIX):
                text = para.text.strip()
                if text:
                    headings.append(text)
        return headings

    def _find_rules_table(self) -> Tuple:
        """
        Locates the rules table and returns the table object and a column index map.
        Raises TemplateParseError if not found.
        """
        if not self.doc.tables:
            raise TemplateParseError("No tables found in the document.")
        
        for table in self.doc.tables:
            if len(table.rows) < 2:
                continue
            header_texts = [cell.text.strip().lower() for cell in table.rows[0].cells]
            if self.REQUIRED_HEADERS.issubset(set(header_texts)):
                col_map = {name: idx for idx, name in enumerate(header_texts)}
                return table, col_map
                
        raise TemplateParseError("Rules table with expected headers ('Section', 'Obligatoire', 'Mots minimum') not found.")

    def extract_rules(self) -> List[ValidationRule]:
        """Parses the 3-column rules table from the document."""
        table, col_map = self._find_rules_table()
        
        section_col = col_map['section']
        mandatory_col = col_map['obligatoire']
        min_words_col = col_map['mots minimum']
        
        rules = []
        for row in table.rows[1:]:
            cells = row.cells
            max_idx = max(col_map.values())
            if len(cells) <= max_idx:
                continue
                
            section_name = cells[section_col].text.strip()
            mandatory_str = cells[mandatory_col].text.strip().lower()
            min_words_str = cells[min_words_col].text.strip()
            
            # Skip completely empty rows
            if not section_name and not mandatory_str and not min_words_str:
                continue
                
            is_mandatory = mandatory_str in ('oui', 'yes', 'true', '1')
            try:
                min_words = int(min_words_str)
            except ValueError:
                min_words = 0
                
            rules.append(ValidationRule(section=section_name, is_mandatory=is_mandatory, min_words=min_words))
        return rules


# SELF-HEAL: Added missing helper function to bridge DocxParser output with BaseLinter expectations
def parse_template_rules(doc_path: Union[str, Path]) -> List[Dict]:
    """Parses a template .docx and returns a list of rule dictionaries."""
    parser = DocxParser(doc_path)
    rules = parser.extract_rules()
    return [
        {"section": r.section, "mandatory": r.is_mandatory, "min_words": r.min_words}
        for r in rules
    ]